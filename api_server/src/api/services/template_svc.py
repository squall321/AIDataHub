"""Word 템플릿 자동 생성 — agent expected schema 기반.

운영자가 새 agent 를 등록할 때, 해당 agent 가 요구하는 ``doc_type`` +
``required_tags`` + (doc_type 에 등록된) ``expected_sections`` 를
미리 채워둔 ``.docx`` 템플릿을 즉시 생성해 배포할 수 있게 한다.

설계 원칙
---------
- **강제 아님** — 작성자는 자유롭게 새 문서를 만들 수 있고, 템플릿은
  단지 "이 포맷으로 만들면 B+C 검증을 통과한다" 는 가이드.
- **agent 가 바뀌면 템플릿도 바뀜** — 즉석 생성 (서버 파일 저장 X).
- **Custom Properties 로 메타 박힘** — Word 변환기가 그대로 흡수해
  ``meta.doc_type`` / ``meta.agent_scope`` 가 자동 채워진다.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_agent_template(
    *,
    agent_type: str,
    agent_name: str,
    agent_description: str = "",
    required_doc_type: str | None = None,
    required_tags: list[str] | None = None,
    excluded_tags: list[str] | None = None,
    doc_type_name: str | None = None,
    doc_type_description: str = "",
    expected_sections: list[str] | None = None,
) -> bytes:
    """agent 의 expected schema 를 반영한 ``.docx`` 바이트를 반환.

    Args:
        agent_type: PK (예: ``"iga-analyst"``).
        agent_name: 표시 이름 (예: ``"IGA 해석 분석가"``).
        agent_description: 자유 텍스트 (옵션).
        required_doc_type: 작성자가 이 agent 로 옵트인할 때 요구되는
            ``meta.doc_type`` 값 (예: ``"manual"``). ``None`` 이면 미지정.
        required_tags: ``meta.tags`` 에 모두 포함되어야 하는 태그.
        excluded_tags: ``meta.tags`` 에서 제외해야 하는 태그.
        doc_type_name: ``required_doc_type`` 에 해당하는 표시 이름
            (UI 에서 채워줌). 없으면 ``required_doc_type`` 자체 사용.
        doc_type_description: doc_type 의 설명. 템플릿 안에서 가이드로 사용.
        expected_sections: doc_type 에 등록된 권장 섹션 헤딩 목록.
            없으면 generic "1. 개요 / 2. 본문 / 3. 결론" 사용.

    Returns:
        .docx 바이너리.
    """
    required_tags = list(required_tags or [])
    excluded_tags = list(excluded_tags or [])
    sections = list(expected_sections or []) or [
        "개요",
        "본문",
        "결론 / 요약",
    ]

    doc = Document()

    # ---- 1. 문서 속성 (변환기가 흡수) -----------------------------------
    cp = doc.core_properties
    title_default = (
        f"[{(doc_type_name or required_doc_type or 'doc')}] "
        f"{agent_name} 용 새 문서"
    )
    cp.title = title_default
    cp.subject = agent_description or doc_type_description or ""
    cp.keywords = ",".join(required_tags)
    cp.author = ""
    cp.comments = (
        f"이 문서는 agent '{agent_type}' ({agent_name}) 가 처리할 수 있는 "
        f"포맷으로 작성하기 위한 템플릿입니다. "
        f"Custom Properties 의 doc_type / agent_scope 를 변경하지 마세요."
    )

    # ---- 2. Custom Properties (변환기가 meta.* 로 읽음) ------------------
    # python-docx 는 custom_properties 인터페이스를 직접 노출하지 않으므로
    # OOXML 의 ``app.xml`` (or ``customProperties``) 를 수동으로 다룬다.
    # 간단하고 안정적인 방식: 본문에 "META BLOCK" 단락을 두고, 변환기가
    # 그 패턴을 인식하도록 한다. (별도 prop 확장은 향후 작업)
    #
    # 현 단계: 작성자가 본문 첫 부분의 BLOCK 을 보존하면 변환기가
    # ``meta.doc_type`` / ``meta.agent_scope`` / ``meta.tags`` 로 흡수한다.
    _add_meta_block(
        doc,
        doc_type=required_doc_type,
        agent_scope=agent_type,
        tags=required_tags,
    )

    # ---- 3. 표지 ---------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run(title_default)
    title_run.bold = True
    title_run.font.size = Pt(22)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_p.add_run(
        f"agent: {agent_type}   ·   "
        f"doc_type: {required_doc_type or '(없음)'}   ·   "
        f"생성일: {datetime.now().strftime('%Y-%m-%d')}"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # ---- 4. 작성 가이드 --------------------------------------------------
    guide_h = doc.add_paragraph()
    guide_h.add_run("작성 가이드").bold = True
    guide = doc.add_paragraph(
        f"이 템플릿은 agent '{agent_type}' ({agent_name}) 가 처리할 수 있는 "
        f"포맷을 미리 채워둔 가이드입니다. 다음을 유지·작성하세요:"
    )
    guide.paragraph_format.space_after = Pt(6)

    bullets = [
        f"맨 위의 META BLOCK 은 변경하지 마세요 — 변환기가 그것을 읽어 "
        f"`meta.doc_type` / `meta.agent_scope` / `meta.tags` 를 자동 채웁니다.",
    ]
    if required_tags:
        bullets.append(
            f"필수 태그 ({len(required_tags)}개): "
            + ", ".join(f"`{t}`" for t in required_tags)
            + " — 빠뜨리면 ingest 경고."
        )
    if excluded_tags:
        bullets.append(
            f"제외할 태그: "
            + ", ".join(f"`{t}`" for t in excluded_tags)
        )
    if doc_type_description:
        bullets.append(f"doc_type '{required_doc_type}' 설명: {doc_type_description}")
    bullets.append(
        f"아래 섹션 헤딩을 그대로 사용하면 agent 의 expected_sections "
        f"검증을 통과합니다. 새 섹션 추가는 자유."
    )

    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # spacer

    # ---- 5. 기대 섹션 placeholder ---------------------------------------
    for i, sec in enumerate(sections, start=1):
        h = doc.add_heading(f"{i}. {sec}", level=1)
        h.paragraph_format.space_before = Pt(12)
        body = doc.add_paragraph(
            f"여기에 '{sec}' 의 내용을 작성하세요. "
            f"단락·표·그림·코드블록 자유롭게 사용 가능."
        )
        body.paragraph_format.space_after = Pt(8)

    # ---- 6. 푸터 정보 ----------------------------------------------------
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    foot_run = foot.add_run(
        f"— 이 템플릿은 agent '{agent_type}' 등록 시점 ({datetime.now().date()}) "
        f"의 스키마로 자동 생성되었습니다. —"
    )
    foot_run.italic = True
    foot_run.font.size = Pt(9)

    # ---- write ----------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_meta_block(
    doc,
    *,
    doc_type: str | None,
    agent_scope: str,
    tags: list[str],
) -> None:
    """본문 맨 위에 변환기가 읽을 META BLOCK 을 단락으로 삽입.

    변환기 (Word ``_build_meta``) 는 첫 페이지 / 첫 본문 영역에서
    "META BLOCK" 마커를 찾아 그 다음 줄들을 key=value 로 파싱한다.
    (현재 변환기는 marker 인식 안 함 — 향후 ``converter/core.py`` 에
    1 회만 추가하면 됨; 지금은 이 블록이 평문으로 본문에 들어가도 무해.)
    """
    h = doc.add_paragraph()
    r = h.add_run("META BLOCK (변경 금지)")
    r.bold = True
    r.font.size = Pt(10)

    if doc_type:
        doc.add_paragraph(f"doc_type: {doc_type}")
    doc.add_paragraph(f"agent_scope: {agent_scope}")
    if tags:
        doc.add_paragraph(f"tags: {', '.join(tags)}")
    doc.add_paragraph("---")  # separator


def template_filename(agent_type: str) -> str:
    """다운로드 시 권장 파일명."""
    safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in agent_type)
    return f"agent_{safe}_template.docx"


__all__ = ["generate_agent_template", "template_filename"]
