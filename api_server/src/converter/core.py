"""Word → JSON 변환 핵심 로직.

설계 원칙:
- Section.blocks 가 본문 등장 순서를 보존한다.
- 표/그림은 본문 흐름 안에 ref 블록으로 삽입되며, 데이터는 최상위 tables/figures 배열에 저장.
- 연속된 코드/등폭 단락은 하나의 code 블록으로 병합한다.
- 이렇게 하면 AI가 blocks를 순서대로 읽으면서 원본 Word 흐름을 그대로 복원할 수 있다.
"""
from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

from .docx_parser import (
    SECTION_NUM_PATTERN,
    coerce_cell_value,
    detect_heading_level,
    extract_attachment_rids,
    extract_image_rids,
    extract_marker,
    extract_section_id_and_title,
    extract_table_data,
    has_inline_image,
    infer_attachment_kind,
    is_caption_paragraph,
    iter_block_items,
    list_marker,
    looks_like_code,
    open_document,
    paragraph_text,
    parse_caption,
)
from .models import (
    Attachment,
    Block,
    ConversionResult,
    Figure,
    Section,
    Source,
    Table,
)

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Migration 0007 — agent-discovery 자동 기본값 헬퍼
# ---------------------------------------------------------------------------

def _default_agent_hints(
    *, data_type_name: str, tags: list[str], section_count: int,
    table_count: int, figure_count: int,
) -> str:
    """본문 통계로 agent_hints 한 줄을 한국어로 생성."""
    topics = ", ".join(t for t in (tags or []) if t) or "N/A"
    return (
        f"이 record 는 {data_type_name} 입니다. 주요 토픽: {topics}. "
        f"본문은 총 {section_count} 섹션, {table_count} 표, "
        f"{figure_count} 그림으로 구성됩니다."
    )


def _default_query_examples(*, title: str, tags: list[str]) -> list[str]:
    """제목 + 태그에서 자연어 질의 예시를 최대 3개 생성."""
    out: list[str] = []
    t = (title or "").strip()
    if t:
        out.append(f"{t} 어떻게 사용해?")
    first_tag = next((x for x in (tags or []) if isinstance(x, str) and x.strip()), None)
    if first_tag:
        out.append(f"{first_tag}에 대해 알려줘")
    if t:
        out.append(f"{t} 관련 자료 보여줘")
    return out[:3]


def _apply_agent_discovery_defaults(
    meta: dict[str, Any],
    *,
    overrides: dict[str, Any] | None = None,
    data_type_name: str,
    title: str,
    tags: list[str],
    section_count: int,
    table_count: int,
    figure_count: int,
) -> None:
    """meta dict 에 Migration 0007 의 4개 필드를 in-place 채운다.

    이미 ``meta`` 에 값이 있거나 ``overrides`` 에서 명시 제공된 경우 그 값을 우선.
    """
    overrides = overrides or {}

    # agent_hints
    if "agent_hints" in overrides and overrides["agent_hints"] is not None:
        meta["agent_hints"] = overrides["agent_hints"]
    elif meta.get("agent_hints") is None:
        meta["agent_hints"] = _default_agent_hints(
            data_type_name=data_type_name,
            tags=tags,
            section_count=section_count,
            table_count=table_count,
            figure_count=figure_count,
        )

    # related_record_ids — 자동 추론 없음, 수동 큐레이션 only.
    if "related_record_ids" in overrides and overrides["related_record_ids"] is not None:
        meta["related_record_ids"] = list(overrides["related_record_ids"])
    elif "related_record_ids" not in meta:
        meta["related_record_ids"] = []

    # query_examples
    if "query_examples" in overrides and overrides["query_examples"] is not None:
        meta["query_examples"] = list(overrides["query_examples"])
    elif not meta.get("query_examples"):
        meta["query_examples"] = _default_query_examples(title=title, tags=tags)

    # access_pattern
    if "access_pattern" in overrides and overrides["access_pattern"]:
        meta["access_pattern"] = overrides["access_pattern"]
    elif not meta.get("access_pattern"):
        meta["access_pattern"] = "occasional"


# ---------------------------------------------------------------------------
# Bug A-3 — summary / tags 자동 추출 (extractive)
#
# 외부 NLP 라이브러리 의존 없이 본문에서 요약 후보 / 키워드 후보를 뽑아낸다.
# CLI 의 --summary / --tags 가 명시되지 않은 문서에서도 RAG hit-rate 에
#필요한 최소한의 신호를 채워주는 게 목적이다.
# ---------------------------------------------------------------------------

import re as _re

_KO_STOPWORDS: set[str] = {
    "및", "이", "그", "저", "것", "수", "등", "또는", "그리고", "하는",
    "하다", "이다", "있다", "없다", "되다", "위해", "통해", "있는", "없는",
    "하지만", "그러나", "더", "또", "또한", "다른", "같은", "이러한",
    "그러한", "어떤", "어떻게", "왜", "누가", "언제", "어디", "무엇", "이런",
}
_EN_STOPWORDS: set[str] = {
    "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
    "of", "with", "by", "from", "as", "is", "are", "was", "were", "be",
    "been", "being", "have", "has", "had", "do", "does", "did", "will",
    "would", "could", "should", "may", "might", "must", "shall", "can",
    "this", "that", "these", "those", "i", "you", "he", "she", "it",
    "we", "they", "them", "their", "our", "my", "your", "his", "her",
    "its", "not", "no", "yes",
}
_ALL_STOPWORDS = _KO_STOPWORDS | _EN_STOPWORDS

# RAKE 분리자 — 한/영 구두점 + 공백.
_PHRASE_SPLIT_RE = _re.compile(
    r"[\s,;:\.!\?\(\)\[\]\{\}\"'`~/\\\|<>=\+\*\&\^%\$#@…·•※→←↑↓"
    r"，。、；：「」『』（）【】《》〈〉]+"
)


def _flatten_paragraph_texts(sections: list[Any]) -> list[str]:
    """sections 트리에서 paragraph/list_item 텍스트만 평탄화 추출."""
    out: list[str] = []

    def walk(nodes: list[Any]) -> None:
        for s in nodes:
            for b in getattr(s, "blocks", []) or []:
                if getattr(b, "type", None) in ("paragraph", "list_item"):
                    txt = getattr(b, "text", None)
                    if txt:
                        out.append(txt)
            walk(getattr(s, "children", []) or [])

    walk(sections)
    return out


def _section_paragraph_texts(section: Any, limit: int) -> list[str]:
    """단일 section 의 첫 N 개 paragraph 블록 텍스트만 반환."""
    out: list[str] = []
    for b in getattr(section, "blocks", []) or []:
        if getattr(b, "type", None) == "paragraph":
            txt = getattr(b, "text", None)
            if txt and txt.strip():
                out.append(txt.strip())
                if len(out) >= limit:
                    break
    return out


def _truncate_at_sentence(text: str, limit: int) -> str:
    """limit 글자 안에서 가장 마지막 문장 종결자에서 자른다.

    ``다.`` (한국어) / ``.`` / ``!`` / ``?`` 어느 것이든 OK.
    경계가 없으면 limit 위치에서 hard-cut.
    """
    if len(text) <= limit:
        return text
    window = text[:limit]
    # 한국어 "다." 우선 — 가장 늦은 위치를 찾는다.
    cut = -1
    for marker in ("다.", "요.", "함.", "음."):
        idx = window.rfind(marker)
        if idx >= 0:
            cut = max(cut, idx + len(marker))
    # 일반 종결자.
    for ch in (".", "!", "?"):
        idx = window.rfind(ch)
        if idx >= 0:
            cut = max(cut, idx + 1)
    if cut > 0 and cut >= int(limit * 0.4):
        return window[:cut].rstrip()
    return window.rstrip()


def _extract_summary_from_sections(
    sections: list[Any], limit: int = 500
) -> str:
    """sections 트리에서 추출형 summary 1 개를 만든다.

    1. section[0] 의 첫 3 개 paragraph 블록을 우선 사용.
    2. 비어 있으면, 다른 섹션의 paragraph 첫 1 개로 fallback.
    3. 공백 1 개로 join → strip → limit 글자에서 문장 경계 컷.
    """
    if not sections:
        return ""
    paras = _section_paragraph_texts(sections[0], limit=3)
    if not paras:
        # fallback: 어느 섹션이든 첫 1 개 paragraph.
        for s in sections:
            paras = _section_paragraph_texts(s, limit=1)
            if paras:
                break
            for child in getattr(s, "children", []) or []:
                paras = _section_paragraph_texts(child, limit=1)
                if paras:
                    break
            if paras:
                break
    if not paras:
        return ""
    joined = " ".join(p.strip() for p in paras if p and p.strip()).strip()
    if not joined:
        return ""
    return _truncate_at_sentence(joined, limit)


def _extract_tags_from_sections(
    sections: list[Any], top_n: int = 8
) -> list[str]:
    """RAKE-style 키워드 추출. 외부 라이브러리 없이 동작.

    ① 본문 paragraph 모두 모아 lowercase.
    ② 구두점/공백으로 phrase 분리, 각 phrase 안에서 stopword 만나면 또 분리.
    ③ phrase = 단어 1~4 개. 각 단어에 대해 freq, deg 계산.
       deg[w] = 그 단어가 등장한 phrase 의 평균 phrase length 합.
       score[phrase] = sum_word(deg[w] / freq[w]).
       (RAKE 표준 변형 — phrase 가 길수록 가산점.)
    ④ 점수 상위 phrase 후보를 뽑은 후 길이/all-stopword/all-numeric 필터.
    ⑤ substring/exact 중복 제거 → top_n 개 반환.
    """
    if not sections:
        return []

    paragraphs = _flatten_paragraph_texts(sections)
    if not paragraphs:
        return []

    # python-rake 가 있으면 우선 사용 — 단, 임포트가 실패하면 fallback.
    try:  # pragma: no cover - optional dep
        import RAKE  # type: ignore
        rake = RAKE.Rake(list(_ALL_STOPWORDS))
        ranked = rake.run(" ".join(paragraphs).lower(), maxWords=4, minFrequency=1)
        return _post_filter_tags([p for p, _s in ranked], top_n)
    except Exception:
        pass

    # ------ minimal RAKE ------
    text = " ".join(paragraphs).lower()
    # phrase 후보 — 비단어 분리자로 split, 그 안에서 stopword 또 분리.
    raw_phrases: list[list[str]] = []
    for chunk in _PHRASE_SPLIT_RE.split(text):
        words = [w for w in chunk.strip().split() if w]
        cur: list[str] = []
        for w in words:
            if w in _ALL_STOPWORDS or w.isdigit():
                if cur:
                    raw_phrases.append(cur)
                    cur = []
                continue
            cur.append(w)
        if cur:
            raw_phrases.append(cur)

    # phrase 길이 1~4 만 채택.
    phrases = [p for p in raw_phrases if 1 <= len(p) <= 4]
    if not phrases:
        return []

    # freq / deg 계산 (RAKE 표준).
    freq: dict[str, int] = {}
    deg: dict[str, int] = {}
    for p in phrases:
        plen = len(p)
        for w in p:
            freq[w] = freq.get(w, 0) + 1
            deg[w] = deg.get(w, 0) + (plen - 1)
    word_score: dict[str, float] = {}
    for w, f in freq.items():
        word_score[w] = (deg[w] + f) / max(f, 1)

    # phrase 점수 = 단어 점수 합. 동일 phrase 는 합쳐서 카운트.
    phrase_score: dict[str, float] = {}
    for p in phrases:
        key = " ".join(p)
        phrase_score[key] = phrase_score.get(key, 0.0) + sum(
            word_score.get(w, 0.0) for w in p
        )

    ranked = sorted(phrase_score.items(), key=lambda kv: kv[1], reverse=True)
    # 길이 2~4 우선, 길이 1 (단어 1개) 도 허용은 하되 후순위.
    multi = [p for p, _s in ranked if 2 <= len(p.split()) <= 4]
    single = [p for p, _s in ranked if len(p.split()) == 1]
    candidates = multi + single
    return _post_filter_tags(candidates, top_n)


def _post_filter_tags(candidates: list[str], top_n: int) -> list[str]:
    """공통 후처리: stopword/all-numeric/너무 짧은 토큰 필터 + dedup."""
    out: list[str] = []
    seen_lower: set[str] = set()
    for raw in candidates:
        if not raw:
            continue
        phrase = raw.strip()
        if not phrase:
            continue
        words = phrase.split()
        if not (1 <= len(words) <= 4):
            continue
        # all-numeric / all-stopword 거르기.
        if all(w.isdigit() for w in words):
            continue
        if all(w in _ALL_STOPWORDS for w in words):
            continue
        # 너무 짧은 1글자 영어 토큰 제거.
        if len(phrase) < 2:
            continue
        low = phrase.lower()
        if low in seen_lower:
            continue
        # substring dedup — 기존 항목이 이번 phrase 를 포함하면 skip.
        if any(low in s or s in low for s in seen_lower):
            continue
        seen_lower.add(low)
        out.append(phrase)
        if len(out) >= top_n:
            break
    return out


@dataclass
class ConverterOptions:
    team: str
    group: str
    year: int
    seq: int = 1
    output_dir: Path = field(default_factory=lambda: Path("output"))
    extract_images: bool = True


@dataclass
class _PendingCaption:
    """그림/표 위치와 매칭 대기 중인 캡션 정보."""

    kind: str
    number: int
    caption_text: str


def _make_doc_id(opts: ConverterOptions) -> str:
    return f"{opts.team}-{opts.group}-{opts.year}-{opts.seq:010d}"


def _make_fig_id(doc_id: str, n: int) -> str:
    return f"{doc_id}-F{n:03d}"


def _make_tbl_id(doc_id: str, n: int) -> str:
    return f"{doc_id}-T{n:03d}"


def _make_att_id(doc_id: str, n: int) -> str:
    return f"{doc_id}-A{n:03d}"


def _sha256_of_file(path: str, chunk: int = 65536) -> str:
    h = hashlib.sha256()
    try:
        with open(path, "rb") as f:
            for buf in iter(lambda: f.read(chunk), b""):
                h.update(buf)
        return h.hexdigest()
    except OSError:
        return ""


class Converter:
    """Word → JSON 변환기."""

    def __init__(self, options: ConverterOptions) -> None:
        self.opts = options
        self.doc_id = _make_doc_id(options)
        self.warnings: list[str] = []

        self.section_root: list[Section] = []
        self.section_stack: list[Section] = []
        self.figures: list[Figure] = []
        self.tables: list[Table] = []
        self.sources: list[Source] = []
        self.attachments: list[Attachment] = []

        self.fig_counter = 0
        self.tbl_counter = 0
        self.src_counter = 0
        self.att_counter = 0

        # auto section numbering
        self.l1_counter = 0
        self.l2_counter = 0
        self.l3_counter = 0

        self.meta_overrides: dict[str, Any] = {}
        self.pending_caption: _PendingCaption | None = None

        # 코드 블록 누적 버퍼 (연속된 등폭 단락을 하나로 묶음)
        self._code_buffer: list[str] = []

        # 휴리스틱 헤딩 감지 결과: paragraph 등장 인덱스 -> (level, sid, title)
        self._heuristic_headings: dict[int, tuple[int, str, str]] = {}

        # 그림 fig_id → drawing relation id 목록 (이미지 추출 단계에서 사용)
        self._figure_rids: dict[str, list[str]] = {}
        # attachment id → list[(rid, hint)] (binary 추출 단계에서 사용)
        self._attachment_rids: dict[str, list[tuple[str, str]]] = {}
        self._document: DocumentType | None = None

    # ---- 외부 API ----

    def convert(self, docx_path: str) -> ConversionResult:
        document = open_document(docx_path)
        self._document = document
        self._heuristic_headings = self._prescan_headings(document)
        # Pre-detect: does the document open with the [DOC_TYPE]/.../[SOURCES]
        # marker prologue convention? If so, suppress the "no Heading 1 at top"
        # warning — the prologue is intentional.
        self._has_marker_prologue = False
        for p in document.paragraphs[:30]:
            t = (p.text or "").strip()
            if t.startswith(("[DOC_TYPE]", "[SUMMARY]", "[TAGS]",
                             "[AGENT_SCOPE]", "[SOURCES]")):
                self._has_marker_prologue = True
                break
        self._process_body(document)
        self._flush_code_buffer()  # 마지막에 남은 코드 블록 정리

        # 그림 바이너리 추출 — 옵션 활성화된 경우만.
        if self.opts.extract_images and self.figures:
            self._extract_figure_binaries(document)

        # 첨부 (figure 일반화) 바이너리 추출 — 모든 종류 포함.
        if self.opts.extract_images:
            self._extract_attachment_binaries(document)

        meta = self._build_meta(document, docx_path)
        return ConversionResult(
            schema_version="1.0",
            meta=meta,
            sections=self.section_root,
            figures=self.figures,
            tables=self.tables,
            sources=self.sources,
            attachments=self.attachments,
            warnings=self.warnings,
        )

    # ---- 내부 처리 ----

    def _process_body(self, document: DocumentType) -> None:
        para_idx = 0
        for el in iter_block_items(document):
            if el.kind == "paragraph" and el.paragraph is not None:
                self._handle_paragraph(el.paragraph, para_idx)
                para_idx += 1
            elif el.kind == "table" and el.table is not None:
                self._flush_code_buffer()
                self._handle_table(el.table)

    def _handle_paragraph(self, p: Paragraph, para_idx: int) -> None:
        text = paragraph_text(p)

        # 마커 처리
        marker = extract_marker(text) if text else None
        if marker:
            self._flush_code_buffer()
            self._handle_marker(marker[0], marker[1])
            return

        # 헤딩 처리 (1순위: Word Heading 스타일)
        level = detect_heading_level(p)
        if level is not None and text:
            self._flush_code_buffer()
            self._open_section(level, text)
            return

        # 헤딩 처리 (2순위: 휴리스틱 감지 — 사전스캔에서 확정된 단락)
        h = self._heuristic_headings.get(para_idx)
        if h is not None:
            h_level, h_sid, h_title = h
            self._flush_code_buffer()
            self._open_section(h_level, f"{h_sid} {h_title}")
            return

        # 캡션 처리
        if text and is_caption_paragraph(p):
            self._flush_code_buffer()
            self._handle_caption(text)
            return

        # 그림 단락 처리
        if has_inline_image(p):
            self._flush_code_buffer()
            self._handle_figure_paragraph(p)
            return

        # 본문 내용 (텍스트 또는 빈 단락)
        # 코드/등폭 판정
        is_code = bool(text) and looks_like_code(p)
        if is_code:
            # 코드 누적 (빈 줄도 포함되어야 하므로 _code_buffer에는 텍스트가 있는 라인만)
            self._code_buffer.append(text)
            return

        # 코드가 아니면, 누적된 코드 블록을 먼저 비운다
        self._flush_code_buffer()

        if not text:
            return

        # 목록 항목 vs 일반 단락
        marker_str = list_marker(p)
        if not self.section_stack:
            self._create_virtual_top_heading()
        cur = self.section_stack[-1]

        if marker_str:
            cur.blocks.append(Block(type="list_item", text=text, marker=marker_str))
        else:
            cur.blocks.append(Block(type="paragraph", text=text))

    def _flush_code_buffer(self) -> None:
        """누적된 코드 블록을 현재 섹션의 blocks에 추가."""
        if not self._code_buffer:
            return
        if not self.section_stack:
            self._create_virtual_top_heading()
        text = "\n".join(self._code_buffer)
        self.section_stack[-1].blocks.append(Block(type="code", text=text))
        self._code_buffer = []

    def _handle_table(self, tbl: Any) -> None:
        self.tbl_counter += 1
        headers, raw_rows = extract_table_data(tbl)
        rows = [[coerce_cell_value(c) for c in row] for row in raw_rows]

        if not headers:
            self.warnings.append(f"표 {self.tbl_counter}: 헤더가 비어 있음")
            headers = [f"col{i + 1}" for i in range(max(len(r) for r in rows) if rows else 1)]

        if rows and any(len(r) != len(headers) for r in rows):
            self.warnings.append(
                f"표 {self.tbl_counter}: 일부 행이 헤더 길이와 다름"
            )

        section_ref = self.section_stack[-1].id if self.section_stack else "0"
        tbl_id = _make_tbl_id(self.doc_id, self.tbl_counter)

        caption_text = self._consume_pending_caption_for("table", self.tbl_counter)
        if caption_text is None:
            caption_text = f"Table {self.tbl_counter}: (캡션 누락 — 검수 필요)"
            self.warnings.append(f"표 {self.tbl_counter}: 캡션 없음")

        t = Table(
            id=tbl_id,
            number=self.tbl_counter,
            caption=caption_text,
            section_ref=section_ref,
            headers=headers,
            rows=rows,
        )
        self.tables.append(t)

        # 본문 흐름에 표 위치 표시
        if not self.section_stack:
            self._create_virtual_top_heading()
        cur = self.section_stack[-1]
        cur.blocks.append(Block(type="table", ref=tbl_id))
        if tbl_id not in cur.table_refs:
            cur.table_refs.append(tbl_id)

    def _handle_caption(self, text: str) -> None:
        parsed = parse_caption(text)
        if not parsed:
            return
        kind, number, full_caption = parsed
        # 직전 그림에 캡션 누락 자동 캡션이 붙어 있으면 즉시 교체
        if (
            kind == "figure"
            and self.figures
            and "(캡션 누락 — 검수 필요)" in self.figures[-1].caption
        ):
            fig_num = self.figures[-1].number
            self.figures[-1].caption = full_caption
            # Attachment 도 동일하게 갱신.
            for att in reversed(self.attachments):
                if "(캡션 누락 — 검수 필요)" in att.caption and att.kind == "figure":
                    att.caption = full_caption
                    break
            # 이전에 기록된 "그림 N: 캡션 없음" 경고를 회수 — 캡션이 해소되었음.
            warn_msg = f"그림 {fig_num}: 캡션 없음"
            self.warnings = [w for w in self.warnings if w != warn_msg]
            return
        # 직전 표에 캡션 누락 자동 캡션이 붙어 있으면 즉시 교체
        if (
            kind == "table"
            and self.tables
            and "(캡션 누락 — 검수 필요)" in self.tables[-1].caption
        ):
            tbl_num = self.tables[-1].number
            self.tables[-1].caption = full_caption
            warn_msg = f"표 {tbl_num}: 캡션 없음"
            self.warnings = [w for w in self.warnings if w != warn_msg]
            return
        # 그렇지 않으면 다음 그림/표를 기다리는 pending 큐에 보관
        self.pending_caption = _PendingCaption(
            kind=kind, number=number, caption_text=full_caption
        )

    def _handle_figure_paragraph(self, p: Paragraph) -> None:
        self.fig_counter += 1
        fig_id = _make_fig_id(self.doc_id, self.fig_counter)
        section_ref = self.section_stack[-1].id if self.section_stack else "0"

        caption_text = self._consume_pending_caption_for("figure", self.fig_counter)
        if caption_text is None:
            caption_text = f"Figure {self.fig_counter}: (캡션 누락 — 검수 필요)"
            self.warnings.append(f"그림 {self.fig_counter}: 캡션 없음")

        fig = Figure(
            id=fig_id,
            number=self.fig_counter,
            caption=caption_text,
            section_ref=section_ref,
        )
        self.figures.append(fig)

        # 그림 단락에서 drawing relation id 를 수집 (나중에 binary 추출에 사용)
        rids = extract_image_rids(p)
        if rids:
            self._figure_rids[fig_id] = rids

        # 첨부 일반화 — 모든 그림은 attachment(kind=figure) 로도 표현된다.
        # 추가로 OLE/오브젝트 rid 도 함께 수집해 동일 단락에서 발견된 비-그림
        # 첨부를 별도 attachment 로 기록한다.
        attachment_rids = extract_attachment_rids(p)
        primary_added = False
        for rid, hint in attachment_rids:
            if hint == "image" and not primary_added:
                # 본 그림 자체를 첨부로 등록.
                self.att_counter += 1
                att_id = _make_att_id(self.doc_id, self.att_counter)
                att = Attachment(
                    id=att_id,
                    number=self.att_counter,
                    kind="figure",
                    caption=caption_text,
                    section_ref=section_ref,
                    extra={"figure_ref": fig_id},
                )
                self.attachments.append(att)
                self._attachment_rids[att_id] = [(rid, hint)]
                primary_added = True
            elif hint == "object":
                # OLE / package 임베디드 object — 별도 첨부.
                self.att_counter += 1
                att_id = _make_att_id(self.doc_id, self.att_counter)
                # caption 은 그림 캡션을 차용 (상세는 _extract_attachment_binaries
                # 단계에서 file_name 기반으로 보강).
                att = Attachment(
                    id=att_id,
                    number=self.att_counter,
                    kind="other",  # 아직 확장자 모름 → 추출 단계에서 갱신
                    caption=caption_text,
                    section_ref=section_ref,
                )
                self.attachments.append(att)
                self._attachment_rids[att_id] = [(rid, hint)]

        if not self.section_stack:
            self._create_virtual_top_heading()
        cur = self.section_stack[-1]
        cur.blocks.append(Block(type="figure", ref=fig_id))
        if fig_id not in cur.figure_refs:
            cur.figure_refs.append(fig_id)

    def _extract_figure_binaries(self, document: DocumentType) -> None:
        """수집된 ``_figure_rids`` 를 따라 docx 의 image part 를 디스크로 복사.

        파일은 ``{output_dir}/{doc_id}/F{nnn}.{ext}`` 에 저장된다.
        해당 ``Figure.image_path`` 에는 ``{doc_id}/F{nnn}.{ext}`` 상대 경로
        문자열을 채운다 (정적 마운트 ``/figures`` 직하 경로와 동일).

        매칭되는 relation 이 없거나 (텍스트 전용 ASCII 다이어그램 등)
        파일 쓰기에 실패하면 ``image_path`` 는 None 으로 남는다.
        """
        if not self._figure_rids:
            return

        out_root = Path(self.opts.output_dir) / self.doc_id
        try:
            out_root.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            self.warnings.append(
                f"그림 출력 폴더 생성 실패: {out_root} ({e})"
            )
            return

        doc_part = document.part
        # rels 는 dict-like (rel_id → relationship)
        rels = doc_part.rels

        for fig in self.figures:
            rids = self._figure_rids.get(fig.id)
            if not rids:
                continue
            saved = False
            for rid in rids:
                rel = rels.get(rid)
                if rel is None:
                    continue
                try:
                    target_part = rel.target_part
                except Exception:  # noqa: BLE001
                    continue
                target_ref = getattr(rel, "target_ref", "") or ""
                ext = self._infer_image_ext(target_ref, target_part)
                file_name = f"F{fig.number:03d}.{ext}"
                out_path = out_root / file_name
                try:
                    blob = target_part.blob
                except Exception as e:  # noqa: BLE001
                    self.warnings.append(
                        f"그림 {fig.number}: 이미지 blob 추출 실패 ({e})"
                    )
                    continue
                try:
                    out_path.write_bytes(blob)
                except OSError as e:
                    self.warnings.append(
                        f"그림 {fig.number}: 이미지 파일 쓰기 실패 — {out_path} ({e})"
                    )
                    continue
                # 상대 경로: '{doc_id}/F001.png' (정적 마운트 /figures 기준)
                fig.image_path = f"{self.doc_id}/{file_name}"
                saved = True
                break  # 한 그림에 대해 첫 매칭 image 만 저장

            if not saved and self.opts.extract_images:
                # 매칭 실패 — 본문에서는 그림이 감지됐지만 image part 를 못 찾은 경우
                self.warnings.append(
                    f"그림 {fig.number}: image part 매칭 실패 (image_path 비어 있음)"
                )

    def _extract_attachment_binaries(self, document: DocumentType) -> None:
        """수집된 ``_attachment_rids`` 를 따라 docx part 를 디스크로 복사.

        파일은 ``{output_dir}/{doc_id}/A{nnn}.{ext}`` 에 저장되며,
        ``Attachment.file_path`` 에는 cross-platform 호환을 위해 항상
        POSIX-style (forward slashes) 상대경로를 채워 넣는다 (예:
        ``"DOC-HE-CAE-2026-0000000001/A001.pdf"``).

        kind 가 아직 ``"other"`` 로 남아있는 첨부는 part 의 확장자/MIME
        으로 다시 추정한다.

        매칭 실패 시 ``file_path`` 는 ``None`` 으로 남고 경고가 추가된다.
        """
        if not self._attachment_rids:
            return

        out_root = Path(self.opts.output_dir) / self.doc_id
        try:
            out_root.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            self.warnings.append(
                f"첨부 출력 폴더 생성 실패: {out_root} ({e})"
            )
            return

        doc_part = document.part
        rels = doc_part.rels

        for att in self.attachments:
            entries = self._attachment_rids.get(att.id)
            if not entries:
                continue

            saved = False
            for rid, _hint in entries:
                rel = rels.get(rid)
                if rel is None:
                    continue
                try:
                    target_part = rel.target_part
                except Exception:  # noqa: BLE001
                    continue
                target_ref = getattr(rel, "target_ref", "") or ""

                # 확장자 결정: target_ref 우선, content_type 폴백.
                ext = self._infer_attachment_ext(target_ref, target_part)
                content_type = (
                    getattr(target_part, "content_type", "") or ""
                ).lower() or None
                file_basename = (
                    Path(target_ref).name if target_ref else f"A{att.number:03d}.{ext}"
                )
                # kind 확정 (이미 figure 면 그대로 둠).
                if att.kind == "other":
                    att.kind = infer_attachment_kind(
                        filename=file_basename, mime=content_type
                    )

                # 출력 파일명: 본래 파일명을 보존하면 확장자 식별/검수가 쉬움.
                # 충돌 방지를 위해 prefix `A{nnn}_` 를 붙인다.
                if target_ref and "." in Path(target_ref).name:
                    safe_name = Path(target_ref).name.replace(" ", "_")
                    file_name = f"A{att.number:03d}_{safe_name}"
                else:
                    file_name = f"A{att.number:03d}.{ext}"

                out_path = out_root / file_name
                try:
                    blob = target_part.blob
                except Exception as e:  # noqa: BLE001
                    self.warnings.append(
                        f"첨부 {att.number}: blob 추출 실패 ({e})"
                    )
                    continue
                try:
                    out_path.write_bytes(blob)
                except OSError as e:
                    self.warnings.append(
                        f"첨부 {att.number}: 파일 쓰기 실패 — {out_path} ({e})"
                    )
                    continue

                # POSIX-style 상대 경로 — DB/API 모두 forward slashes 로 통일.
                rel_path = Path(self.doc_id) / file_name
                att.file_path = rel_path.as_posix()
                att.file_name = file_name
                if content_type:
                    att.mime_type = content_type
                try:
                    att.size_bytes = out_path.stat().st_size
                except OSError:
                    att.size_bytes = None
                saved = True
                break

            if not saved:
                self.warnings.append(
                    f"첨부 {att.number}: relation part 매칭 실패 (file_path 비어 있음)"
                )

    @staticmethod
    def _infer_attachment_ext(target_ref: str, target_part: Any) -> str:
        """첨부 part 의 확장자 추정 (이미지 외에도 사용)."""
        if target_ref:
            tail = target_ref.rsplit(".", 1)[-1].lower()
            if tail and tail.replace("_", "").isalnum() and len(tail) <= 8:
                return tail
        ct = (getattr(target_part, "content_type", "") or "").lower()
        if "/" in ct:
            sub = ct.split("/", 1)[1]
            sub = sub.split(";", 1)[0].strip()
            if sub:
                return "jpg" if sub == "jpeg" else sub
        return "bin"

    @staticmethod
    def _infer_image_ext(target_ref: str, target_part: Any) -> str:
        """target_ref 또는 part content_type 으로 확장자 추정. 기본 png."""
        if target_ref:
            tail = target_ref.rsplit(".", 1)[-1].lower()
            if tail and tail.isalnum() and len(tail) <= 5:
                return tail
        ct = (getattr(target_part, "content_type", "") or "").lower()
        # "image/png" → "png", "image/jpeg" → "jpeg"
        if "/" in ct:
            sub = ct.split("/", 1)[1]
            sub = sub.split(";", 1)[0].strip()
            if sub:
                return "jpg" if sub == "jpeg" else sub
        return "png"

    def _consume_pending_caption_for(self, kind: str, number: int) -> str | None:
        pc = self.pending_caption
        if pc is None or pc.kind != kind:
            return None
        self.pending_caption = None
        return pc.caption_text

    def _handle_marker(self, key: str, value: str) -> None:
        if key == "DOC_TYPE":
            self.meta_overrides["doc_type"] = value
        elif key == "SUMMARY":
            self.meta_overrides["summary"] = value
        elif key == "TAGS":
            self.meta_overrides["tags"] = [t.strip() for t in value.split(",") if t.strip()]
        elif key == "AGENT_SCOPE":
            self.meta_overrides["agent_scope"] = [
                t.strip() for t in value.split(",") if t.strip()
            ]
        elif key == "SOURCES":
            # The SOURCES marker is followed by a Table Grid (parsed as a normal
            # table). We don't extract its rows specially in this version, but
            # the table itself is preserved in `tables[]`. This is by design
            # (KooRemapper prologue convention) — no warning emitted.
            self.meta_overrides["has_sources_marker"] = True

    def _open_section(self, level: int, heading_text: str) -> None:
        parsed_id, title = extract_section_id_and_title(heading_text)
        auto_id = self._next_auto_id(level)
        section_id = parsed_id or auto_id
        # Mismatch between author-supplied numbering and auto-numbering is
        # **normal** for documents with explicit chapter numbers (e.g., the
        # Theory manual: "1.2 Solid Elements" / "23.4 Material Models"). The
        # author's number is canonical; emit a warning only when the depth
        # disagrees (which would indicate a real structural problem).
        if parsed_id and parsed_id != auto_id:
            depth_parsed = parsed_id.count(".") + 1
            depth_auto = auto_id.count(".") + 1
            if depth_parsed != depth_auto:
                self.warnings.append(
                    f"섹션 번호 불일치 (깊이): 본문='{parsed_id}', "
                    f"자동={auto_id}. 본문 값 사용."
                )

        section = Section(id=section_id, level=level, title=title)

        while self.section_stack and self.section_stack[-1].level >= level:
            self.section_stack.pop()

        if not self.section_stack:
            self.section_root.append(section)
        else:
            self.section_stack[-1].children.append(section)

        self.section_stack.append(section)

    def _next_auto_id(self, level: int) -> str:
        if level == 1:
            self.l1_counter += 1
            self.l2_counter = 0
            self.l3_counter = 0
            return f"{self.l1_counter}"
        if level == 2:
            self.l2_counter += 1
            self.l3_counter = 0
            return f"{self.l1_counter}.{self.l2_counter}"
        self.l3_counter += 1
        return f"{self.l1_counter}.{self.l2_counter}.{self.l3_counter}"

    def _create_virtual_top_heading(self) -> None:
        # If the doc uses the marker prologue convention (DOC_TYPE/SUMMARY/...
        # before any Heading 1), creating a virtual '본문' section is the
        # designed behaviour — not a problem worth warning about.
        if not getattr(self, "_has_marker_prologue", False):
            self.warnings.append("문서 시작에 Heading 1 없음 → 가상 '본문' 섹션 추가")
        s = Section(id="1", level=1, title="본문")
        self.section_root.append(s)
        self.section_stack.append(s)
        self.l1_counter = 1

    def _prescan_headings(
        self, document: DocumentType
    ) -> dict[int, tuple[int, str, str]]:
        """Heading 스타일이 적용되지 않은 단락 중 번호 패턴이 있는 것을 헤딩으로 추정.

        2-pass 알고리즘:
        - level 2/3 후보 (예: 1.1, 2.3.1)는 항상 헤딩으로 확정
        - level 1 후보 (예: 1., 2.)는 그 prefix에 해당하는 level 2 후보가
          뒤에 등장하는 경우에만 헤딩으로 확정 (단순 번호 목록과 구별)
        - 같은 prefix를 가진 level 1 후보가 여럿이면 첫 level 2 직전 것을 선택

        키는 paragraph 등장 순서 인덱스 (lxml 프록시 재사용 회피).

        Returns:
            dict {paragraph_index: (level, section_id, title)}
        """
        # main pass와 동일한 순회 순서로 paragraph 인덱스 부여
        # iter_block_items는 paragraph + table을 모두 yield하므로 paragraph만 카운팅
        candidates: list[tuple[int, int, str, str]] = []
        # (paragraph_index, level, section_id, title)

        para_idx = 0
        for el in iter_block_items(document):
            if el.kind != "paragraph" or el.paragraph is None:
                continue
            p = el.paragraph

            # Heading 스타일이 이미 있으면 건너뜀
            if detect_heading_level(p) is None:
                text = paragraph_text(p).strip()
                if text:
                    m = SECTION_NUM_PATTERN.match(text)
                    if m:
                        sid = m.group(1)
                        title = m.group(2).strip()
                        level = sid.count(".") + 1
                        if 1 <= level <= 3:
                            candidates.append((para_idx, level, sid, title))

            para_idx += 1

        if not candidates:
            return {}

        confirmed: dict[int, tuple[int, str, str]] = {}

        # level 2/3 후보는 모두 확정
        for idx, level, sid, title in candidates:
            if level >= 2:
                confirmed[idx] = (level, sid, title)

        # 각 prefix별 첫 번째 level 2 등장 위치 수집
        first_l2_idx: dict[str, int] = {}
        for idx, level, sid, _title in candidates:
            if level == 2:
                prefix = sid.split(".")[0]
                if prefix not in first_l2_idx or idx < first_l2_idx[prefix]:
                    first_l2_idx[prefix] = idx

        # 각 prefix에 대해 그 직전(가장 가까운 앞)의 level 1 후보를 헤딩으로 확정
        for prefix, l2_idx in first_l2_idx.items():
            l1_before = [
                (idx, sid, title)
                for idx, level, sid, title in candidates
                if level == 1 and sid == prefix and idx < l2_idx
            ]
            if l1_before:
                best = max(l1_before, key=lambda c: c[0])
                confirmed[best[0]] = (1, best[1], best[2])

        # 연속 시퀀스 규칙:
        # 마지막으로 confirmed된 level 1 (예: 10) 다음에 등장하는
        # 11, 12, 13... 후보는 sub-heading 없어도 자동 확정 (장 끝부분 처리).
        confirmed_l1 = sorted(
            [(idx, sid) for idx, (lvl, sid, _t) in confirmed.items() if lvl == 1],
            key=lambda c: c[0],
        )
        if confirmed_l1:
            last_idx, last_sid = confirmed_l1[-1]
            try:
                expected_num = int(last_sid) + 1
                for idx, level, sid, title in candidates:
                    if level != 1 or idx <= last_idx or idx in confirmed:
                        continue
                    try:
                        num = int(sid)
                    except ValueError:
                        continue
                    if num == expected_num:
                        confirmed[idx] = (1, sid, title)
                        last_idx = idx
                        expected_num = num + 1
            except ValueError:
                pass

        if confirmed:
            self.warnings.append(
                f"Heading 스타일 미적용 → 휴리스틱으로 {len(confirmed)}개 단락을 "
                f"헤딩으로 추정함. 정확한 결과를 위해 Word [제목 1/2/3] 스타일 권장."
            )

        return confirmed

    def _build_meta(self, document: DocumentType, docx_path: str) -> dict[str, Any]:
        core = document.core_properties
        title = core.title or Path(docx_path).stem
        author = core.author or ""
        created = (core.created or datetime.now(tz=timezone.utc)).strftime("%Y-%m-%d")
        modified = (core.modified or datetime.now(tz=timezone.utc)).strftime("%Y-%m-%d")

        # Bug A-3: --summary / --tags 가 빠진 docx 라도 본문에서 추출.
        summary_override = (self.meta_overrides.get("summary") or "").strip()
        if summary_override:
            summary = summary_override
        else:
            summary = _extract_summary_from_sections(self.section_root)

        tags_override = list(self.meta_overrides.get("tags") or [])
        auto_tags = _extract_tags_from_sections(self.section_root, top_n=8)
        merged_tags = list(dict.fromkeys(tags_override + auto_tags))[:20]

        meta: dict[str, Any] = {
            "doc_id": self.doc_id,
            "title": title,
            "source_format": "docx",
            "source_file": Path(docx_path).name,
            "doc_type": self.meta_overrides.get("doc_type", "manual"),
            "created": created,
            "modified": modified,
            "author": author,
            "department": f"{self.opts.team}-{self.opts.group}",
            "version": "1.0",
            "tags": merged_tags,
            "summary": summary,
        }
        if "agent_scope" in self.meta_overrides:
            meta["agent_scope"] = self.meta_overrides["agent_scope"]

        # Migration 0007: agent-discovery 자동 기본값.
        section_count = len(self.section_root)
        table_count = len(self.tables)
        figure_count = len(self.figures)
        _apply_agent_discovery_defaults(
            meta,
            overrides=self.meta_overrides,
            data_type_name="Word 문서",
            title=str(title),
            tags=meta["tags"],
            section_count=section_count,
            table_count=table_count,
            figure_count=figure_count,
        )

        if not meta["tags"]:
            self.warnings.append("[TAGS] 마커 없음 → tags 비어 있음")
        if not meta["summary"]:
            self.warnings.append("[SUMMARY] 마커 없음 → summary 비어 있음")

        return meta


def write_output(
    result: ConversionResult,
    output_dir: Path,
) -> tuple[Path, Path]:
    """JSON과 경고 로그 저장."""
    import json

    output_dir.mkdir(parents=True, exist_ok=True)

    doc_id = result.meta["doc_id"]
    json_path = output_dir / f"{doc_id}.json"
    log_path = output_dir / f"{doc_id}.warnings.log"

    json_path.write_text(
        json.dumps(result.to_dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    if result.warnings:
        log_path.write_text(
            "\n".join(f"[WARN] {w}" for w in result.warnings),
            encoding="utf-8",
        )
    elif log_path.exists():
        log_path.unlink()

    return json_path, log_path


__all__ = [
    "Converter",
    "ConverterOptions",
    "write_output",
]
