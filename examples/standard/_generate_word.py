"""표준 Word(.docx) 예제 생성 스크립트.

생성되는 sample_report.docx 가 시연하는 원칙
(word_to_json_conversion_rules.md):
- Heading 1/2 스타일 사용 (Heading 감지 규칙)
- 표 캡션은 표 위 (캡션 감지 규칙)
- [TAGS] / [SUMMARY] / [AGENT_SCOPE] 마커 단락 (자동 생성 필드)
- 각 섹션마다 실제 본문 단락 1~2개

CAE/IGA 도메인 테마 — iga_guide.docx 와 일관.

실행 방법
--------
python _generate_word.py [출력경로]
기본 출력: 같은 폴더의 sample_report.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def build_sample_report(out_path: Path) -> Path:
    doc = Document()

    # ── 머리말: 메타 마커 단락 (Word 변환 규칙서 부록 A) ──
    doc.add_paragraph("[DOC_TYPE] report")
    doc.add_paragraph(
        "[SUMMARY] KooRemapper IGA 변환 입력 시료의 사양과 변환 결과를 "
        "정리한 표준 예제 보고서. 각 변환 규칙(헤딩·표 캡션·메타 마커)을 "
        "최소 단위로 시연한다."
    )
    doc.add_paragraph("[TAGS] IGA, NURBS, KooRemapper, sample, standard")
    doc.add_paragraph("[AGENT_SCOPE] iga-analyst, cae-reporter")

    # ── 1. 개요 ──
    doc.add_heading("1. 개요", level=1)

    doc.add_heading("1.1 목적", level=2)
    doc.add_paragraph(
        "본 문서는 표준 Word 작성 규칙(Heading 스타일, 표 캡션 위, "
        "그림 캡션 아래)을 그대로 따르는 최소 예제이다. 사용자는 이 파일을 "
        "그대로 복사하고 내용만 교체해 작성을 시작할 수 있다."
    )

    doc.add_heading("1.2 범위", level=2)
    doc.add_paragraph(
        "본 예제는 IGA 변환에 사용되는 시료 2종(알루미늄 6061, 일반 강재)의 "
        "재질·치수만 다룬다. 실제 변환 결과·검증 곡선 등은 별도 부속 문서로 "
        "분리한다."
    )

    # ── 2. 시료 사양 ──
    doc.add_heading("2. 시료 사양", level=1)
    doc.add_paragraph(
        "다음 표는 변환 검증에 사용된 두 시료의 재질과 두께를 요약한다. "
        "표 캡션은 표 위에 배치한다(Word 작성 3원칙 중 하나)."
    )

    # 표 캡션 — 표 위 (Word 변환 규칙 6.3절 — 표 캡션은 표 위)
    cap = doc.add_paragraph()
    cap.style = doc.styles["Caption"]
    cap.add_run("Table 1: 시료 사양 (재질·치수)")

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    headers = ["시료ID", "재질", "두께(mm)"]
    rows = [
        ["S001", "Al6061", "2.0"],
        ["S002", "Steel", "1.5"],
    ]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, r in enumerate(rows, start=1):
        for j, v in enumerate(r):
            table.rows[i].cells[j].text = v

    # ── 3. 결론 ──
    doc.add_heading("3. 결론", level=1)
    doc.add_paragraph(
        "두 시료 모두 KooRemapper IGA 변환 파이프라인의 입력 요건"
        "(닫힌 FE mesh, 단일 PART)을 충족하였다. 변환 결과 *.iga 파일은 "
        "재해석 검증을 통과하였다."
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def main() -> int:
    out = (
        Path(sys.argv[1])
        if len(sys.argv) > 1
        else Path(__file__).resolve().parent / "sample_report.docx"
    )
    p = build_sample_report(out)
    print(f"[OK] Word 예제 생성: {p}  ({p.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
